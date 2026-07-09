from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

PURPLE = RGBColor(0x5B, 0x21, 0xB6)
DARK = RGBColor(0x1E, 0x29, 0x3B)
GRAY = RGBColor(0x64, 0x74, 0x8B)
BLUE = RGBColor(0x25, 0x63, 0xEB)
GREEN = RGBColor(0x05, 0x96, 0x69)
ORANGE = RGBColor(0xEA, 0x58, 0x0C)

def add_h1(text, color=PURPLE):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = color

def add_h2(text, color=DARK):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = color
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)

def add_h3(text, color=BLUE):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = color
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)

def add_meta(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = GRAY

def add_para(text, bold=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    if bold:
        r.bold = True
    return p

def add_label_value(label, value):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r2 = p.add_run(value)

def add_message_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    for line in text.split("\n"):
        r = p.add_run(line + "\n")
        r.font.name = 'Consolas'
        r.font.size = Pt(10.5)
        r.font.color.rgb = DARK

def add_divider():
    p = doc.add_paragraph()
    r = p.add_run("_" * 80)
    r.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

# COVER
add_h1("HeyReach LinkedIn Campaigns", PURPLE)
add_para("P0 Passive Pipeline Re-engagement - Xoxoday Plum (Global API)", size=12)
add_meta("Sender account: Gaurav Sava (LinkedIn ID: 168813) | 3 campaigns | April 2026")
add_divider()

# SHARED CADENCE OVERVIEW
add_h2("Cadence Overview (Applied to All 3 Campaigns)")
add_para("11-day cadence, 5 LinkedIn touches:", size=11)
add_para("  Day 1 - Profile Visit (automatic)")
add_para("  Day 2 - Like Recent Post")
add_para("  Day 4 - Connection Request (NO note)")
add_para("  Day 7 - First DM (if connected) or InMail (if not)")
add_para("  Day 10 - Follow-up DM with value-add asset")

add_h2("Merge Tags HeyReach Supports")
add_para("  {{first_name}} - prospect's first name")
add_para("  {{last_name}} - prospect's last name")
add_para("  {{company_name}} - prospect's company")
add_para("  {{position}} - prospect's job title")

add_h2("Campaign Settings (Recommended)")
add_para("  Daily limits: 20 connection requests / 50 messages per day per account")
add_para("  Working hours: 9am-6pm in Gaurav's timezone")
add_para("  Days: Mon-Fri only")
add_para("  Skip if already connected: Yes (start from DM step)")
add_para("  Stop on reply: Yes")

doc.add_page_break()

# ============ CAMPAIGN 1 ============
add_h1("Campaign 1 - P0 Plum API Warm", GREEN)
add_label_value("HeyReach List", "P0-C1-Plum-API-Warm-Apr2026 (ID: 614694)")
add_label_value("Leads", "254 (re-engagement, prospects who previously explored Plum)")
add_label_value("Angle", "We've shipped a lot since you last looked")
add_divider()

add_h3("Day 1 - Profile Visit")
add_para("Automatic profile visit via HeyReach. No message.")

add_h3("Day 2 - Like Recent Post")
add_para("Like prospect's most recent LinkedIn post. Genuine engagement, no comment needed unless substantive.")

add_h3("Day 4 - Connection Request (NO note)")
add_para("Clean connection request. Leave the note field EMPTY. Acceptance rates are higher without a pitch.")

add_h3("Day 7 - First DM (if connected) or InMail (if not)")
add_label_value("Subject (InMail only)", "Following up on Plum, {{first_name}}")
add_para("Message:", bold=True)
add_message_block("""Hey {{first_name}},

I've sent a couple of emails about Plum's rewards API. Not sure if they landed.

Quick version: we power the redemption layer for apps like Viwell and HDB. 10,000+ rewards, 100+ countries, one API.

You explored a Plum integration with {{company_name}} a while back. We've shipped quite a bit since then (Perks API, better SDKs, SOC 2 + ISO 27001 + GDPR).

If it's still on your radar, happy to share sandbox access. If not, no worries at all.

Gaurav""")

add_h3("Day 10 - Follow-up DM / Value Add")
add_para("Message:", bold=True)
add_message_block("""{{first_name}} - thought this might be useful.

Quick breakdown of how Plum compares to Tremendous / Tango / Rybbon (most teams evaluating rewards infra look at all four):

Plum: gift cards + experiences + merchandise + charity + mobile top-ups, all in 100+ countries, one API.
Tremendous: gift cards only.
Tango: primarily US.
Rybbon: email-based delivery, not API-first.

If rewards comes back on {{company_name}}'s roadmap, we're easy to reach. Cheers!

Gaurav""")

doc.add_page_break()

# ============ CAMPAIGN 2 ============
add_h1("Campaign 2 - P0 Passive New Deals", BLUE)
add_label_value("HeyReach List", "P0-C2-Passive-New-Deals-Apr2026 (ID: 614695)")
add_label_value("Leads", "293 (cold intro, generic New Deal entries that went cold)")
add_label_value("Angle", "Earn engine + burn engine, fresh value hook")
add_divider()

add_h3("Day 1 - Profile Visit")
add_para("Automatic profile visit.")

add_h3("Day 2 - Like Recent Post")
add_para("Like prospect's most recent LinkedIn post.")

add_h3("Day 4 - Connection Request (NO note)")
add_para("Clean connection request, no message attached.")

add_h3("Day 7 - First DM (if connected) or InMail (if not)")
add_label_value("Subject (InMail only)", "Quick question about rewards at {{company_name}}")
add_para("Message:", bold=True)
add_message_block("""Hey {{first_name}},

Sent a couple of emails about Plum's rewards API. Thought I'd try here too.

Quick version: if {{company_name}} runs any kind of points, cashback, referral, or engagement program, we power the redemption layer. 10,000+ rewards in 100+ countries, one API, 2-3 weeks to go live.

Companies like Viwell, HDB Financial, and Curefit use it for everything from wellness rewards to banking loyalty.

If it's relevant, happy to share sandbox access. If not, no worries!

Gaurav""")

add_h3("Day 10 - Follow-up DM / Value Add")
add_para("Message:", bold=True)
add_message_block("""{{first_name}} - one concrete example:

A wellness app client was struggling with low redemption rates. Users earning points but not converting because the catalogue was limited. After integrating Plum, they expanded to 10,000+ options across 100+ countries. Redemption rates tripled in 90 days.

If {{company_name}} ever runs into a similar problem, we're a quick conversation away.

Gaurav""")

doc.add_page_break()

# ============ CAMPAIGN 3 ============
add_h1("Campaign 3 - P0 Empuls / Loyalty / Other", ORANGE)
add_label_value("HeyReach List", "P0-C3-Empuls-Loyalty-Other-Apr2026 (ID: 614696)")
add_label_value("Leads", "70 (multi-product angle, Empuls / Loyalty / Incentives / Perks)")
add_label_value("Angle", "Broader Xoxoday platform, works across employees, customers, partners, panelists")
add_divider()

add_h3("Day 1 - Profile Visit")
add_para("Automatic profile visit.")

add_h3("Day 2 - Like Recent Post")
add_para("Like prospect's most recent LinkedIn post.")

add_h3("Day 4 - Connection Request (NO note)")
add_para("Clean connection request.")

add_h3("Day 7 - First DM (if connected) or InMail (if not)")
add_label_value("Subject (InMail only)", "The rewards problem no one wants to build, {{first_name}}")
add_para("Message:", bold=True)
add_message_block("""Hey {{first_name}},

Sent a couple of emails about Xoxoday's rewards platform.

Whether it's employee R&R, customer loyalty, sales incentives, or survey rewards, the pattern's the same: companies build the program logic in-house but don't want to build the catalogue, fulfilment, and compliance layer.

One integration gives {{company_name}} access to 10,000+ rewards in 100+ countries. Works for employees, customers, partners, panelists, whoever you're trying to reward.

If any of those motions are in play, happy to tailor a walkthrough. If not, no worries!

Gaurav""")

add_h3("Day 10 - Follow-up DM / Value Add")
add_para("Message:", bold=True)
add_message_block("""{{first_name}} - three quick examples of what teams do with the same API:

1. Wellness app (Viwell): users earn points for healthy habits, redeem on Plum's catalogue. 3x redemption lift.

2. Neo bank (HDB Financial): credit card rewards, multi-currency fulfilment with SOC 2 + ISO 27001. Faster go-live than building in-house.

3. Market research firm (Nielsen): survey panel incentives in 40+ countries, replaced manual gift card procurement. 60% less ops overhead.

Same infrastructure, three very different outcomes. If {{company_name}}'s ever in a similar spot, we're easy to reach.

Gaurav""")

doc.add_page_break()

# ============ FINAL PAGE ============
add_h1("Summary", PURPLE)
add_para("Total leads across 3 HeyReach campaigns: 617", bold=True)
add_para("  Campaign 1 (Plum API Warm): 254 leads")
add_para("  Campaign 2 (Passive New Deals): 293 leads")
add_para("  Campaign 3 (Empuls / Loyalty / Other): 70 leads")

add_h2("Setup Checklist")
add_para("1. In HeyReach, create 3 new campaigns (one per segment above)")
add_para("2. Attach the pre-loaded lead list to each campaign using the List IDs")
add_para("3. Configure the 5 cadence steps using the copy above")
add_para("4. Assign sender account: Gaurav Sava (ID: 168813)")
add_para("5. Set daily limits, working hours, and reply stop settings")
add_para("6. Preview 3-5 prospects to verify merge tags render correctly")
add_para("7. Launch when ready, coordinate Day 1 with the Smartlead Day 1 email send")

add_meta("Generated April 15, 2026 - Xoxoday ABM | Global API | Plum")

# Save
output_path = "/Users/naitikchavda/Event Auto push/smartlead-kit/outputs/HeyReach_LinkedIn_Copy_All_3_Campaigns.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
